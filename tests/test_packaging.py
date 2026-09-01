"""Packaging / import regression guards."""

from __future__ import annotations

import subprocess
import sys

import openpyxl

from prisma_s import __version__
from prisma_s.keywords import bundled_dict_path, bundled_dict_text, load_keywords
from prisma_s.runner import run_analysis
from prisma_s.search import build_regex


def test_version_is_unified():
    assert __version__ == "1.5.0"


def test_import_prisma_s_does_not_import_pandas():
    code = "import prisma_s, sys; print('pandas' in sys.modules)"
    out = subprocess.run(
        [sys.executable, "-c", code], capture_output=True, text=True, check=True
    ).stdout.strip()
    assert out == "False"


def test_bundled_dicts_exist():
    assert bundled_dict_path().name == "keyword_dictionary_v1.3.csv"
    assert bundled_dict_path().exists()
    assert bundled_dict_path("keyword_dictionary_v1.1.csv").exists()


def test_default_bundled_is_registry():
    kw = load_keywords(None)
    assert kw.is_registry and kw.version == "1.3" and kw.n_canonical == 98


def test_bundled_v13_first_line():
    assert bundled_dict_text().splitlines()[0].split(",")[:3] == [
        "category", "term_id", "canonical_term"
    ]


def test_accented_term_matches_case_insensitively():
    assert len(build_regex("ação").findall("Uma Ação e outra ação")) == 2
    assert build_regex("análisis").findall("El ANÁLISIS final") == ["ANÁLISIS"]


def test_run_analysis_smoke_registry(tmp_path):
    docs = tmp_path / "docs"
    docs.mkdir()
    from docx import Document

    d = Document()
    d.add_paragraph("The supplier and the mill. Supplier again. jurisdiction, landscape, farm.")
    d.save(str(docs / "sample.docx"))

    out = tmp_path / "out" / "smoke.xlsx"
    df = run_analysis(batch_id="smoke", output_xlsx=out, input_path=docs, emit_citation=False)

    wb = openpyxl.load_workbook(out)
    assert wb.sheetnames == [
        "Long_AllTerms", "Term_Summary", "D1_Key_Terms", "Zero_Reference_Terms",
        "PRISMA-S_Compliance", "Run_Metadata",
    ]
    supplier = df[df["Canonical Term"] == "supplier"]["Count"]
    assert int(supplier.iloc[0]) == 2
    assert set(df.columns) >= {"Canonical Term", "Variants Included", "Referenced"}


def test_flat_mode_still_available(tmp_path):
    docs = tmp_path / "docs"
    docs.mkdir()
    from docx import Document

    d = Document()
    d.add_paragraph("Cocoa and Coffee and Coffee.")
    d.save(str(docs / "s.docx"))
    out = tmp_path / "out11.xlsx"
    df = run_analysis(
        batch_id="v11", output_xlsx=out, input_path=docs,
        keyword_csv="bundled:1.1", emit_citation=False, figures=False,
    )
    wb = openpyxl.load_workbook(out)
    assert wb.sheetnames == ["Long_AllTerms", "PRISMA-S_Compliance", "Run_Metadata"]
    assert int(df[df["Term"] == "Coffee"]["Count"].iloc[0]) == 2
